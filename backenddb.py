import io
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse,StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import mysql.connector
from mysql.connector import Error
from pydantic import BaseModel
from passlib.context import CryptContext 

# import io
# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
import docx2txt
from PyPDF2 import PdfReader
import chardet
import langid
from io import BytesIO

# from fastapi import FastAPI, File, UploadFile, HTTPException, Form
# from fastapi.responses import StreamingResponse
# from pydantic import BaseModel
import docx
import requests
from io import BytesIO
from transformers import pipeline
# import torch
from transformers import AutoTokenizer, TFMarianMTModel
from typing import List
from fastapi.middleware.cors import CORSMiddleware


app = FastAPI()
host = "cdaserver.mysql.database.azure.com"
user = "cdaadmin"
password = "fillpassword"
db = "doctrans"

def get_db_connection():
    try:
        connection = mysql.connector.connect(host=host, user=user, password=password, database=db)
        return connection
    except Error as e:
        print("Error while connecting to MySQL", e)
        raise HTTPException(status_code=500, detail="Database connection failed")

# Allow CORS for the frontend
origins = [
    "http://127.0.0.1:5500",
    "http://localhost:5500",
    "http://127.0.0.1:5501",
    "https://immune-grouper-marginally.ngrok-free.app",
    
    
    # Add other allowed origins here
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class RegisterRequest(BaseModel):
    username: str
    email: str
    password: str

# Password hashing context
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def convertBinarytoFile(binarydata, filename):
    file = io.BytesIO(binarydata)
    file.seek(0)  # Ensure the stream is at the beginning
    return file
        

@app.get("/download/translated/{translation_id}")
async def download_translated_document(translation_id: int):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # Query to retrieve the translation based on translation_id
        query = """SELECT t.translated_content, d.document_name, t.language 
                   FROM Translations t 
                   LEFT JOIN Documents d ON t.original_document_id = d.document_id 
                   WHERE t.translation_id = %s"""
        cursor.execute(query, (translation_id,))
        result = cursor.fetchone()

        if result:
            translated_content, original_name, language = result

             # Convert the binary content to a file-like object
            file_content = convertBinarytoFile(translated_content)
            translated_filename = f"{original_name}_{language[:2].lower()}.pdf"

            # Use StreamingResponse to serve the file
            response = StreamingResponse(file_content, media_type='application/pdf')
            response.headers['Content-Disposition'] = f'attachment; filename="{translated_filename}"'
            return response
        else:
            raise HTTPException(status_code=404, detail="Translation not found.")

    except mysql.connector.Error as e:
        print("Failed to retrieve translation.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve translation.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
        
# Registration endpoint
@app.post("/register/")
async def register(request: RegisterRequest):
    connection = get_db_connection()
    cursor = connection.cursor()

    # Hash the user's password
    hashed_password = pwd_context.hash(request.password)

    try:
        insert_user_query = """
        INSERT INTO Users (username, password, email) 
        VALUES (%s, %s, %s)
        """
        cursor.execute(insert_user_query, (request.username, hashed_password, request.email))
        connection.commit()
        return {"message": "User registered successfully"}

    except Error as e:
        connection.rollback()
        if e.errno == 1062:  # Duplicate entry error
            raise HTTPException(status_code=400, detail="Username already exists.")
        else:
            raise HTTPException(status_code=500, detail="Internal server error.")
    
    finally:
        cursor.close()
        connection.close()
        
# Function to store document in the database
def store_document(user_id, document_name, document_content):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        insert_doc_query = """
        INSERT INTO Documents (user_id, document_name, document_content) 
        VALUES (%s, %s, %s)
        """
        cursor.execute(insert_doc_query, (user_id, document_name, document_content))
        connection.commit()
        print(cursor.lastrowid)
        return cursor.lastrowid
    
    except Error as e:
        print("Failed to store document.", e)
        raise HTTPException(status_code=500, detail="Failed to store document.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()



def store_translation(original_document_id, language, translated_content):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # First, execute the SELECT query
        select_document_query = "SELECT * FROM Documents WHERE document_id = %s"
        cursor.execute(select_document_query, (original_document_id,))
        
        # Fetch all the results (even if you don't need them)
        cursor.fetchall()

        # Now insert the translation into the Translations table
        insert_translation_query = """
        INSERT INTO Translations (original_document_id, language, translated_content) 
        VALUES (%s, %s, %s)
        """
        cursor.execute(insert_translation_query, (original_document_id, language, translated_content))
        connection.commit()
       
        return cursor.lastrowid

    except mysql.connector.Error as e:
        print("Failed to store translation.", e)
        raise HTTPException(status_code=500, detail="Failed to store translation.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()



class ForgotPasswordRequest(BaseModel):
    email: str
  

class LoginRequest(BaseModel):
    username: str
    password: str




@app.post("/forgot-password/")
async def forgot_password(request: ForgotPasswordRequest):
    connection = get_db_connection()
    cursor = connection.cursor(buffered=True)  # Use a buffered cursor

    try:
        # Check if the email exists in the Users table
        query = "SELECT user_id FROM Users WHERE email = %s"
        cursor.execute(query, (request.email,))
        result = cursor.fetchone()

        if result is None:
            raise HTTPException(status_code=404, detail="Email not found")

        # Implement your password reset logic here, e.g., generate a reset token and send an email
        # For now, we will just return a success message
        return {"message": "Password reset instructions have been sent to your email."}

    except Error as e:
        raise HTTPException(status_code=500, detail="Failed to process request.")
    
    finally:
        cursor.close()
        connection.close()

# Login endpoint
@app.post("/login/")
async def login(request: LoginRequest):
    connection = get_db_connection()
    cursor = connection.cursor()

    try:
        query = "SELECT password FROM Users WHERE username = %s"
        cursor.execute(query, (request.username,))
        result = cursor.fetchone()
        print(result)

        if result is None or not pwd_context.verify(request.password, result[0]):
            raise HTTPException(status_code=401, detail="Invalid username or password")

        return {"message": "Login successful"}

    finally:
        cursor.close()
        connection.close()


def convertFileToBinary(file):
    binarydata = file.read()
    return binarydata

# Endpoint to handle document upload and translation

@app.post("/upload-original-document/")
async def uploadoriginaldocument(user_id: int = Form(...), language: str = Form(...), file: UploadFile = File(...)):
    # file_content = await file.read()
    binary_content = convertFileToBinary(file.file)
    document_id = store_document(user_id, file.filename, binary_content)

    return JSONResponse(content={"document_id": document_id})




@app.post("/upload-translated-document/")
async def uploadtranslateddocument(doc_id: int = Form(...), language: str = Form(...), file: UploadFile = File(...)):
    file_content = await file.read()  # Read the file content here
    binary_content = convertFileToBinary2(file_content)  # Pass the file content, not the filename
    document_id = store_translation(doc_id, file.filename, binary_content)
    print("done")
    return JSONResponse(content={"document_id": document_id})

def convertFileToBinary2(file_content):
    return file_content  # You don't need file.read() here, just return the content as is

# import base64
@app.get("/alldocuments/")
def get_documents_and_translations():
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # Query to retrieve document details along with their starred status
        select_documents_query = """
        SELECT d.document_name, d.uploaded_at, d.starred, t.language, t.translated_content, t.original_document_id, t.translation_id
        FROM Translations t
        LEFT JOIN Documents d ON t.original_document_id = d.document_id
        """
        cursor.execute(select_documents_query)
        translations = cursor.fetchall()

        if translations:
            document_dict = {}
            for translation in translations:
                doc_name, uploaded_at, starred, lang, trans_content, original_id, translated_id = translation
                if doc_name not in document_dict:
                    document_dict[doc_name] = {
                        "name": doc_name,
                        "translated": [],
                        "sourceLang": "English",
                        "translatedLang": [],
                        "uploadedDate": uploaded_at.strftime("%Y-%m-%d"),
                        "starred": bool(starred),  # Get the starred status from the DB
                        "docID": original_id,
                        "transID": translated_id
                    }
                if lang:
                    translated_doc_name = f"{lang}"
                    document_dict[doc_name]["translated"].append(translated_doc_name)
                    # language = lang.split('_')[1].split('.docx')[0]
                    language = lang.rsplit('_', 1)[-1].split('.docx')[0]
                    print(language)
                    document_dict[doc_name]["translatedLang"].append(language)

            # Format the data to match the frontend structure
            documents_list = []
            for doc_name, details in document_dict.items():
                documents_list.append({
                    "name": details["name"],
                    "translated": ", ".join(details["translated"]),
                    "sourceLang": details["sourceLang"],
                    "translatedLang": ", ".join(details["translatedLang"]),
                    "uploadedDate": details["uploadedDate"],
                    "starred": details["starred"],
                    "docID": details["docID"],
                    "transID": details["transID"]
                })
            return documents_list
        else:
            raise HTTPException(status_code=404, detail="No translations found.")

    except Error as e:
        print("Failed to retrieve translations.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve translations.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
            

    
@app.get("/stareddocuments/")
def get_starred_documents_and_translations():
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # Query to retrieve starred translations and their corresponding original document names and upload dates
        select_documents_query = """
        SELECT d.document_name, d.uploaded_at, t.language, t.translated_content, t.original_document_id, t.translation_id, d.starred
        FROM Translations t
        LEFT JOIN Documents d ON t.original_document_id = d.document_id
        WHERE d.starred = TRUE
        """
        cursor.execute(select_documents_query)
        translations = cursor.fetchall()

        if translations:
            document_dict = {}
            for translation in translations:
                doc_name, uploaded_at, lang, trans_content, original_id, translated_id, starred = translation
                if doc_name not in document_dict:
                    document_dict[doc_name] = {
                        "name": doc_name,
                        "translated": [],
                        "sourceLang": "English",
                        "translatedLang": [],
                        "uploadedDate": uploaded_at.strftime("%Y-%m-%d"),
                        "starred": starred,
                        "docID": original_id,
                        "transID": translated_id
                    }
                if lang:
                    translated_doc_name = f"{lang}"
                    document_dict[doc_name]["translated"].append(translated_doc_name)
                    filename = lang
                    language = filename.split('_')[1].split('.')[0]
                    document_dict[doc_name]["translatedLang"].append(language)

            # Format the data to match the frontend structure
            documents_list = []
            for doc_name, details in document_dict.items():
                documents_list.append({
                    "name": details["name"],
                    "translated": ", ".join(details["translated"]),
                    "sourceLang": details["sourceLang"],
                    "translatedLang": ", ".join(details["translatedLang"]),
                    "uploadedDate": details["uploadedDate"],
                    "starred": details["starred"],
                    "docID": details["docID"],
                    "transID": details["transID"]
                })

            return documents_list
        else:
            raise HTTPException(status_code=404, detail="No starred documents found.")

    except Error as e:
        print("Failed to retrieve translations.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve translations.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()


# Function to download original document
@app.get("/download/original/{document_id}")
async def download_original_document(document_id: int):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # Query to retrieve the document based on document_id
        query = "SELECT document_name, document_content FROM Documents WHERE document_id = %s"
        cursor.execute(query, (document_id,))
        result = cursor.fetchone()

        if result:
            document_name, document_content = result

            # Convert binary content to bytes and prepare it for download
            file_content = io.BytesIO(document_content)

            return FileResponse(file_content, media_type='application/octet-stream', filename=document_name)
        else:
            raise HTTPException(status_code=404, detail="Document not found.")

    except Error as e:
        print("Failed to retrieve document.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve document.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
            


@app.post("/documents/toggle_star/{document_id}")
async def toggle_star(document_id: int):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()

        # Retrieve the current starred status
        query = "SELECT starred FROM Documents WHERE document_id = %s"
        cursor.execute(query, (document_id,))
        result = cursor.fetchone()

        if result is None:
            raise HTTPException(status_code=404, detail="Document not found.")

        # Toggle the starred status
        current_status = result[0]
        new_status = not current_status

        # Update the document's starred status in the database
        update_query = "UPDATE Documents SET starred = %s WHERE document_id = %s"
        cursor.execute(update_query, (new_status, document_id))
        connection.commit()

        return {"starred": new_status}

    except Error as e:
        print("Failed to toggle star status.", e)
        raise HTTPException(status_code=500, detail="Failed to toggle star status.")
    
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
            
 
class TextPayload(BaseModel):
    text: str
 
# Mapping of language codes to full names
language_map = {
    "af": "Afrikaans",
    "am": "Amharic",
    "an": "Aragonese",
    "ar": "Arabic",
    "as": "Assamese",
    "az": "Azerbaijani",
    "be": "Belarusian",
    "bg": "Bulgarian",
    "bn": "Bengali",
    "br": "Breton",
    "bs": "Bosnian",
    "ca": "Catalan",
    "cs": "Czech",
    "cy": "Welsh",
    "da": "Danish",
    "de": "German",
    "dz": "Dzongkha",
    "el": "Greek",
    "en": "English",
    "eo": "Esperanto",
    "es": "Spanish",
    "et": "Estonian",
    "eu": "Basque",
    "fa": "Persian",
    "fi": "Finnish",
    "fo": "Faroese",
    "fr": "French",
    "ga": "Irish",
    "gl": "Galician",
    "gu": "Gujarati",
    "he": "Hebrew",
    "hi": "Hindi",
    "hr": "Croatian",
    "ht": "Haitian Creole",
    "hu": "Hungarian",
    "hy": "Armenian",
    "id": "Indonesian",
    "is": "Icelandic",
    "it": "Italian",
    "ja": "Japanese",
    "jv": "Javanese",
    "ka": "Georgian",
    "kk": "Kazakh",
    "km": "Khmer",
    "kn": "Kannada",
    "ko": "Korean",
    "ku": "Kurdish",
    "ky": "Kyrgyz",
    "la": "Latin",
    "lb": "Luxembourgish",
    "lo": "Lao",
    "lt": "Lithuanian",
    "lv": "Latvian",
    "mg": "Malagasy",
    "mk": "Macedonian",
    "ml": "Malayalam",
    "mn": "Mongolian",
    "mr": "Marathi",
    "ms": "Malay",
    "mt": "Maltese",
    "nb": "Norwegian Bokmål",
    "ne": "Nepali",
    "nl": "Dutch",
    "nn": "Norwegian Nynorsk",
    "no": "Norwegian",
    "oc": "Occitan",
    "or": "Oriya",
    "pa": "Punjabi",
    "pl": "Polish",
    "ps": "Pashto",
    "pt": "Portuguese",
    "qu": "Quechua",
    "ro": "Romanian",
    "ru": "Russian",
    "rw": "Kinyarwanda",
    "se": "Northern Sami",
    "si": "Sinhala",
    "sk": "Slovak",
    "sl": "Slovenian",
    "sq": "Albanian",
    "sr": "Serbian",
    "sv": "Swedish",
    "sw": "Swahili",
    "ta": "Tamil",
    "te": "Telugu",
    "th": "Thai",
    "tl": "Tagalog",
    "tr": "Turkish",
    "ug": "Uighur",
    "uk": "Ukrainian",
    "ur": "Urdu",
    "vi": "Vietnamese",
    "vo": "Volapük",
    "wa": "Walloon",
    "xh": "Xhosa",
    "zh": "Chinese",
    "zu": "Zulu"
}

def extract_text_from_file(file: UploadFile) -> str:
    try:
        if file.content_type == "text/plain":
            content = file.file.read()
            result = chardet.detect(content)
            text = content.decode(result['encoding'])
        elif file.content_type == "application/pdf":
            pdf_reader = PdfReader(BytesIO(file.file.read()))
            text = "".join([page.extract_text() for page in pdf_reader.pages])
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = docx2txt.process(BytesIO(file.file.read()))
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")

@app.post("/detect-language/")
async def detect_language(file: UploadFile = File(...)):
    try:
        text = extract_text_from_file(file)
        if not text:
            raise HTTPException(status_code=400, detail="No text extracted from the file")

        language_code, confidence = langid.classify(text)
        if language_code in language_map:
            return {"language": language_map[language_code]}
        else:
            raise HTTPException(status_code=400, detail="No language detected")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    
src = "en"  # source language
 
 
class TranslationRequest(BaseModel):
    target_lang: str
 
def translate_text(text_chunk,target_lang):
 
    # batch = tokenizer([text_chunk], return_tensors="tf")
    # gen = model.generate(**batch)
    # t =tokenizer.batch_decode(gen, skip_special_tokens=True)
    # print(t[0])
    # return t[0]
    return "text"
 
@app.post("/translate-document/")
async def translate_document(
    target_lang: str = Form(...),
    file: UploadFile = File(...)
):
    try:
        # Read the file
        file_content = await file.read()
        doc = docx.Document(BytesIO(file_content))
        trg = next((code for code, language in language_map.items() if language.lower() == target_lang.lower()), None)
 
        # model_name = f"Helsinki-NLP/opus-mt-{src}-{trg}"
 
        # global model
        # model = TFMarianMTModel.from_pretrained(model_name)
        # global tokenizer
        # tokenizer = AutoTokenizer.from_pretrained(model_name)
 
        # Helper function to translate and style paragraphs
        def translate_and_style_paragraph(paragraph, original_para):
            original_text = original_para.text
            if original_text.strip():
                translated_text = translate_text(original_text, target_lang)
 
                # Preserve list formatting
                if original_para.style.name.startswith('List'):
                    translated_text = translated_text
 
                # Extract and apply styles
                if original_para.runs:
                    first_run = original_para.runs[0]
                    para.clear()
                    new_run = para.add_run(translated_text)
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                    new_run.bold = first_run.bold
                    new_run.italic = first_run.italic
                    new_run.underline = first_run.underline
                    if first_run.font.color:
                        new_run.font.color.rgb = first_run.font.color.rgb
                else:
                    para.text = translated_text
 
        # Translate paragraphs
        for para in doc.paragraphs:
            translate_and_style_paragraph(para, para)
 
        # Translate table contents
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        translate_and_style_paragraph(para, para)
 
        # Translate headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                translate_and_style_paragraph(para, para)
            for para in section.footer.paragraphs:
                translate_and_style_paragraph(para, para)
 
        # Save the updated document to a BytesIO object
        updated_doc_io = BytesIO()
        doc.save(updated_doc_io)
        updated_doc_io.seek(0)
 
        # Return the document as a downloadable file
        headers = {
            'Content-Disposition': f'attachment; filename="translated_{file.filename}"',
            'Content-Type': file.content_type
        }
 
        return StreamingResponse(updated_doc_io, headers=headers)
 
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    
    
